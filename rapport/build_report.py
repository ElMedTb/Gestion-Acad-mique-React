from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "rapport" / "Rapport_Gestion_Academique.docx"
SCREENSHOTS = Path(__file__).resolve().parent / "screenshots"


GREEN = "0B6B3A"
LIGHT_GREEN = "EAF5EF"
LIGHT_GRAY = "F3F5F2"
DARK = "1F2933"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Bloc", "Fonctionnalites", "Statut"]
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], GREEN)
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for bloc, details, status in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], bloc, bold=True)
        set_cell_text(cells[1], details)
        set_cell_text(cells[2], status)
    doc.add_paragraph()


def add_accounts_table(doc):
    add_heading(doc, "5. Comptes de demonstration", 1)
    doc.add_paragraph("Mot de passe commun du jeu de donnees : password123")
    rows = [
        ("ADMIN", "admin@academic.com", "Gestion globale des donnees et des comptes."),
        ("SCOLARITE", "scolarite1@academic.com", "Gestion administrative, etudiants, matieres et notes."),
        ("TEACHER", "teacher1@academic.com", "Matieres IA et Big Data."),
        ("TEACHER", "teacher2@academic.com", "Matieres Cloud et Securite."),
        ("STUDENT", "student1@academic.com", "Notes MBDS avec cas de note eliminatoire."),
        ("STUDENT", "student2@academic.com", "Etudiant MBDS avec double diplomation IA."),
        ("STUDENT", "student3@academic.com", "Etudiant BI avec notes en UE decisionnelle."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(["Role", "Compte", "Utilisation"]):
        shade_cell(table.rows[0].cells[i], GREEN)
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for role, account, usage in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], role, bold=True)
        set_cell_text(cells[1], account)
        set_cell_text(cells[2], usage)
    doc.add_paragraph()


def add_screenshot_slot(doc, number, title, filename, caption):
    add_heading(doc, f"Capture {number} - {title}", 3)
    path = SCREENSHOTS / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(8)
    else:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        shade_cell(cell, LIGHT_GRAY)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Capture a inserer : rapport/screenshots/{filename}\n{caption}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("53606A")
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Rapport de projet\nGestion academique")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Application React / Node.js - gestion des etudiants, matieres, UE, diplomes, professeurs, notes et notifications")
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor.from_string("53606A")

    doc.add_paragraph()
    add_heading(doc, "1. Presentation generale", 1)
    doc.add_paragraph(
        "Ce rapport presente une application web de gestion academique developpee avec React pour le frontend "
        "et Node.js/Express pour le backend. Elle centralise les donnees academiques, securise l'acces par role "
        "et propose des parcours adaptes aux profils ADMIN, SCOLARITE, STUDENT et TEACHER."
    )

    add_heading(doc, "2. Architecture technique", 1)
    add_bullets(doc, [
        "Frontend React/Vite avec Material UI, theme clair/sombre et interface responsive.",
        "Backend Node.js/Express avec MongoDB et Mongoose.",
        "Authentification JWT, SSO Google et OTP optionnel.",
        "Notifications plateforme, email et SMS avec gestion des destinataires.",
        "Dockerfile frontend, Dockerfile backend et docker-compose.yml pour lancer l'ensemble.",
    ])

    add_heading(doc, "3. Couverture fonctionnelle", 1)
    add_status_table(doc, [
        ("Modele academique", "Etudiants, matieres, UE, diplomes, professeurs, notes et double diplomation.", "OK"),
        ("Authentification", "Connexion classique, SSO Google, OTP optionnel et roles exacts.", "OK"),
        ("Droits d'acces", "ADMIN complet, SCOLARITE administratif, STUDENT personnel, TEACHER limite a ses matieres.", "OK"),
        ("Notifications", "Ciblage utilisateur, historique, email de compte, email/SMS selon le cas.", "OK"),
        ("Statistiques", "Dashboards par profil, moyennes, progression, classement et exports.", "OK"),
        ("Containerisation", "Lancement complet via Docker Compose avec MongoDB, backend et frontend.", "OK"),
    ])

    add_heading(doc, "4. Scenarios principaux", 1)
    scenarios = [
        ("Connexion administrateur", "L'ADMIN se connecte, ouvre le dashboard et gere les donnees globales."),
        ("Creation d'un etudiant", "La scolarite renseigne nom, prenom, email, photo, telephone, niveau, filiere et double diplomation."),
        ("Saisie des notes", "La note est affectee a une matiere, visible par UE et retrouvee dans l'espace etudiant."),
        ("Connexion etudiant Google", "L'etudiant utilise son email Google et consulte uniquement ses notes et statistiques."),
        ("Espace professeur", "Le TEACHER voit uniquement les UE/matieres qui lui sont affectees et peut saisir les notes correspondantes."),
        ("Notification ciblee", "Un destinataire est selectionne, recoit la notification plateforme et, selon le canal, email ou SMS."),
    ]
    for name, detail in scenarios:
        p = doc.add_paragraph(style="List Number")
        p.add_run(name + " : ").bold = True
        p.add_run(detail)

    add_accounts_table(doc)

    add_heading(doc, "6. Captures et preuves", 1)
    slots = [
        ("1", "Connexion", "01_connexion.png", "Page de connexion avec le formulaire, le champ OTP et le bouton Google."),
        ("2", "Dashboard ADMIN", "02_admin_dashboard.png", "Vue globale des entites academiques et statistiques principales."),
        ("3", "Creation etudiant", "03_admin_etudiants.png", "Formulaire etudiant avec photo, telephone, niveau, filiere et double diplomation."),
        ("4", "Notes SCOLARITE", "04_scolarite_notes.png", "Saisie des notes et association etudiant-cours."),
        ("5", "Espace TEACHER", "05_teacher_espace.png", "Vue limitee aux UE et matieres affectees au professeur."),
        ("6", "Espace STUDENT", "06_student_notes.png", "Notes personnelles, moyennes et progression de l'etudiant."),
        ("7", "Barre de notifications", "07_notifications_barre.png", "Dernieres notifications et compteur lu/non lu."),
        ("8", "Historique notifications", "08_notifications_historique.png", "Liste des notifications ciblees et generales."),
        ("9", "Creation notification", "09_notification_creation.png", "Choix du destinataire et des canaux plateforme/email/SMS."),
        ("10", "Email creation compte", "13_email_compte.jpg", "Email contenant les informations de connexion temporaires."),
        ("11", "Email notification", "14_email_notification.jpeg", "Email recu apres creation d'une notification importante."),
        ("12", "SMS notification", "15_sms_notification.jpg", "SMS recu sur telephone apres envoi via Twilio."),
        ("13", "Docker", "11_docker.png", "Lancement de l'application complete avec docker compose."),
    ]
    for slot in slots:
        add_screenshot_slot(doc, *slot)

    add_heading(doc, "7. Verification", 1)
    add_bullets(doc, [
        "Backend : cd backend puis npm run test:smoke.",
        "Frontend : cd frontend puis npm run build.",
        "Docker : docker compose up --build.",
        "Notifications : creer un compte, verifier l'email d'identifiants, creer une notification ciblee et verifier plateforme/email/SMS.",
    ])

    add_heading(doc, "8. Conclusion", 1)
    doc.add_paragraph(
        "Le projet repond aux criteres principaux : gestion academique complete, authentification et roles, "
        "persistance des donnees, statistiques, notifications, documentation, collection Postman et containerisation. "
        "Les captures jointes permettent de demontrer les parcours essentiels de chaque role."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
